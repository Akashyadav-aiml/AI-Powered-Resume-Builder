from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Depends
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
from PyPDF2 import PdfReader
from docx import Document
import io
import re

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import llm_helper as llm_ops
from auth import create_access_token, decode_token, verify_password, get_password_hash, Token

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

app = FastAPI()
api_router = APIRouter(prefix="/api")

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Auth Security
security = HTTPBearer()

# Models - Auth
class UserRegister(BaseModel):
    email: EmailStr
    full_name: str
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    id: str
    email: str
    full_name: str
    created_at: datetime

# Models
class ResumeSection(BaseModel):
    section_name: str
    content: str

class ResumeData(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    raw_text: str
    sections: List[ResumeSection]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ATSScore(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    resume_id: str
    overall_score: int
    keyword_score: int
    formatting_score: int
    section_score: int
    details: Dict[str, Any]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class EnhancedResume(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    original_resume_id: str
    enhanced_text: str
    enhanced_sections: List[ResumeSection]
    enhancement_type: str
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ManualResumeInput(BaseModel):
    full_name: str
    email: str
    phone: str
    summary: str
    experience: str
    education: str
    skills: str

class EnhanceRequest(BaseModel):
    resume_id: str
    enhancement_type: str = "both"

class GoogleAuthRequest(BaseModel):
    credential: str  # Google ID token

# Helper Functions
def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        raise HTTPException(status_code=400, detail="Failed to extract text from PDF")

def extract_text_from_docx(file_content: bytes) -> str:
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        raise HTTPException(status_code=400, detail="Failed to extract text from DOCX")

def parse_resume_sections(text: str) -> List[ResumeSection]:
    sections = []
    section_patterns = [
        (r"(SUMMARY|PROFESSIONAL SUMMARY|OBJECTIVE)\s*:?\s*([\s\S]*?)(?=\n\n|EXPERIENCE|EDUCATION|SKILLS|$)", "Summary"),
        (r"(EXPERIENCE|WORK EXPERIENCE|EMPLOYMENT)\s*:?\s*([\s\S]*?)(?=\n\n|EDUCATION|SKILLS|$)", "Experience"),
        (r"(EDUCATION|ACADEMIC BACKGROUND)\s*:?\s*([\s\S]*?)(?=\n\n|EXPERIENCE|SKILLS|$)", "Education"),
        (r"(SKILLS|TECHNICAL SKILLS|COMPETENCIES)\s*:?\s*([\s\S]*?)(?=\n\n|EXPERIENCE|EDUCATION|$)", "Skills"),
    ]
    
    for pattern, section_name in section_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            sections.append(ResumeSection(section_name=section_name, content=match.group(2).strip()))
    
    if not sections:
        sections.append(ResumeSection(section_name="Content", content=text))
    
    return sections

def calculate_ats_score(text: str, sections: List[ResumeSection]) -> ATSScore:
    keyword_list = [
        "python", "javascript", "react", "node", "fastapi", "mongodb", "sql",
        "aws", "docker", "kubernetes", "git", "agile", "scrum", "ci/cd",
        "leadership", "communication", "teamwork", "problem-solving",
        "bachelor", "master", "degree", "certified", "manager", "engineer"
    ]
    
    text_lower = text.lower()
    keyword_matches = sum(1 for kw in keyword_list if kw in text_lower)
    keyword_score = min(100, int((keyword_matches / len(keyword_list)) * 100))
    
    section_names = [s.section_name.lower() for s in sections]
    required_sections = ["experience", "education", "skills"]
    section_matches = sum(1 for req in required_sections if any(req in s for s in section_names))
    section_score = int((section_matches / len(required_sections)) * 100)
    
    has_email = bool(re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text))
    has_phone = bool(re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text))
    has_consistent_format = len(text.split('\n')) > 5
    formatting_score = int(((has_email + has_phone + has_consistent_format) / 3) * 100)
    
    overall_score = int((keyword_score * 0.4 + section_score * 0.4 + formatting_score * 0.2))
    
    return ATSScore(
        resume_id="",
        overall_score=overall_score,
        keyword_score=keyword_score,
        formatting_score=formatting_score,
        section_score=section_score,
        details={
            "keyword_matches": keyword_matches,
            "total_keywords": len(keyword_list),
            "has_email": has_email,
            "has_phone": has_phone,
            "sections_found": [s.section_name for s in sections]
        }
    )

async def enhance_with_openai(text: str) -> str:
    return await llm_ops.enhance_with_openai(text)

async def enhance_with_gemini(text: str) -> str:
    return await llm_ops.enhance_with_gemini(text)

def generate_pdf(resume_data: dict) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=0.6*inch, leftMargin=0.6*inch,
        topMargin=0.55*inch, bottomMargin=0.45*inch
    )

    BLACK  = colors.HexColor('#1A1A1A')
    DGRAY  = colors.HexColor('#2D2D2D')
    MGRAY  = colors.HexColor('#555555')
    ACCENT = colors.HexColor('#1E3A5F')

    name_style = ParagraphStyle('PDFName', fontName='Helvetica-Bold', fontSize=22,
                                 textColor=BLACK, alignment=TA_CENTER, spaceAfter=3, spaceBefore=0)
    contact_style = ParagraphStyle('PDFContact', fontName='Helvetica', fontSize=8.5,
                                    textColor=MGRAY, alignment=TA_CENTER, spaceAfter=8)
    sec_style = ParagraphStyle('PDFSec', fontName='Helvetica-Bold', fontSize=10,
                                textColor=ACCENT, spaceBefore=10, spaceAfter=1)
    entry_title_style = ParagraphStyle('PDFETitle', fontName='Helvetica-Bold', fontSize=9.5,
                                        textColor=BLACK, spaceBefore=5, spaceAfter=0)
    entry_sub_style   = ParagraphStyle('PDFESub', fontName='Helvetica-Oblique', fontSize=9,
                                        textColor=MGRAY, spaceBefore=0, spaceAfter=2)
    entry_date_style  = ParagraphStyle('PDFEDate', fontName='Helvetica', fontSize=9,
                                        textColor=MGRAY, alignment=TA_RIGHT, spaceBefore=5, spaceAfter=0)
    entry_date_sub    = ParagraphStyle('PDFEDateSub', fontName='Helvetica', fontSize=9,
                                        textColor=MGRAY, alignment=TA_RIGHT, spaceBefore=0, spaceAfter=2)
    bullet_style = ParagraphStyle('PDFBullet', fontName='Helvetica', fontSize=9.5,
                                   textColor=DGRAY, leftIndent=14, spaceAfter=2, leading=13)
    plain_style  = ParagraphStyle('PDFPlain', fontName='Helvetica', fontSize=9.5,
                                   textColor=DGRAY, spaceAfter=3, leading=14)

    DATE_RE = re.compile(
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,.]+\d{4}'
        r'|\bpresent\b|\d{4}\s*[-\u2013]\s*(?:\d{4}|present)',
        re.IGNORECASE
    )

    def _parse_entries(content):
        entries = []
        cur_h, cur_b = [], []
        for raw in content.split('\n'):
            s = raw.strip()
            if not s:
                continue
            if s[0] in '-\u2022*\u00b7\u2013':
                cur_b.append(s.lstrip('-\u2022*\u00b7\u2013 ').strip())
            else:
                if cur_b:
                    entries.append((cur_h, cur_b))
                    cur_h, cur_b = [s], []
                else:
                    cur_h.append(s)
        if cur_h or cur_b:
            entries.append((cur_h, cur_b))
        return entries

    def _split_date(line):
        m = DATE_RE.search(line)
        if m:
            date_str = line[m.start():].strip()
            left = line[:m.start()].strip().rstrip('|,\u2013- ').strip()
            return left, date_str
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            return parts[0], ' | '.join(parts[1:])
        return line, ''

    def section_block(title):
        return [
            Spacer(1, 4),
            Paragraph(title.upper(), sec_style),
            HRFlowable(width='100%', thickness=0.8, color=ACCENT, spaceAfter=3),
        ]

    def render_entry(headers, bullets):
        items = []
        for i, h in enumerate(headers):
            left, right = _split_date(h)
            st = entry_title_style if i == 0 else entry_sub_style
            dt = entry_date_style  if i == 0 else entry_date_sub
            if right:
                row = Table(
                    [[Paragraph(left, st), Paragraph(right, dt)]],
                    colWidths=['72%', '28%']
                )
                row.setStyle(TableStyle([
                    ('VALIGN',        (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING',   (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING',  (0, 0), (-1, -1), 0),
                    ('TOPPADDING',    (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ]))
                items.append(row)
            else:
                items.append(Paragraph(h, st))
        for b in bullets:
            items.append(Paragraph(f'\u2022 {b}', bullet_style))
        return items

    story = []

    # ── Header ──────────────────────────────────────────────────
    full_name = resume_data.get('full_name', resume_data.get('id', 'Resume'))
    story.append(Paragraph(full_name, name_style))

    contact_parts = [p for p in [
        resume_data.get('email', ''), resume_data.get('phone', '')
    ] if p]
    if contact_parts:
        story.append(Paragraph('  |  '.join(contact_parts), contact_style))

    story.append(HRFlowable(width='100%', thickness=1.5, color=ACCENT, spaceAfter=6))

    # ── Sections ────────────────────────────────────────────────
    for section in resume_data.get('sections', []):
        sec_name = section['section_name']
        content  = section.get('content', '')
        story   += section_block(sec_name)

        if 'skill' in sec_name.lower():
            # Render skills as clean inline text
            skills_text = '  •  '.join(
                line.strip() for line in content.split('\n') if line.strip()
            )
            story.append(Paragraph(skills_text, plain_style))
            continue

        entries = _parse_entries(content)
        if not entries:
            story.append(Paragraph(content, plain_style))
            continue

        for headers, bullets in entries:
            story += render_entry(headers, bullets)

    doc.build(story)
    result = buffer.getvalue()
    buffer.close()
    return result

def generate_docx(resume_data: dict) -> bytes:
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.55)
        sec.bottom_margin = Inches(0.45)
        sec.left_margin   = Inches(0.6)
        sec.right_margin  = Inches(0.6)

    # Reuse the default empty first paragraph for the name
    name_para = doc.paragraphs[0]

    ACCENT = RGBColor(0x1E, 0x3A, 0x5F)
    BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
    MGRAY  = RGBColor(0x55, 0x55, 0x55)

    DATE_RE = re.compile(
        r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,.]+\d{4}'
        r'|\bpresent\b|\d{4}\s*[-\u2013]\s*(?:\d{4}|present)',
        re.IGNORECASE
    )

    def sp(para, before=0, after=2):
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after  = Pt(after)

    def add_hr(color='1E3A5F', sz='8'):
        p = doc.add_paragraph()
        sp(p, 0, 1)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), sz)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def set_right_tab(para, pos_twips=10512):
        """Right-aligned tab stop so name\tdate aligns to right margin."""
        pPr = para._p.get_or_add_pPr()
        tabs_el = OxmlElement('w:tabs')
        tab_el  = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:pos'), str(pos_twips))
        tabs_el.append(tab_el)
        pPr.append(tabs_el)

    def _parse_entries(content):
        entries = []
        cur_h, cur_b = [], []
        for raw in content.split('\n'):
            s = raw.strip()
            if not s:
                continue
            if s[0] in '-\u2022*\u00b7\u2013':
                cur_b.append(s.lstrip('-\u2022*\u00b7\u2013 ').strip())
            else:
                if cur_b:
                    entries.append((cur_h, cur_b))
                    cur_h, cur_b = [s], []
                else:
                    cur_h.append(s)
        if cur_h or cur_b:
            entries.append((cur_h, cur_b))
        return entries

    def _split_date(line):
        m = DATE_RE.search(line)
        if m:
            date_str = line[m.start():].strip()
            left = line[:m.start()].strip().rstrip('|,\u2013- ').strip()
            return left, date_str
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            return parts[0], ' | '.join(parts[1:])
        return line, ''

    # ── Name ────────────────────────────────────────────────────
    full_name = resume_data.get('full_name', resume_data.get('id', 'Resume'))
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(name_para, 0, 2)
    nr = name_para.add_run(full_name)
    nr.bold = True
    nr.font.size = Pt(22)
    nr.font.color.rgb = BLACK

    # ── Contact ─────────────────────────────────────────────────
    contact_parts = [p for p in [resume_data.get('email', ''), resume_data.get('phone', '')] if p]
    if contact_parts:
        cp = doc.add_paragraph(' | '.join(contact_parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(cp, 0, 4)
        for r in cp.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = MGRAY

    add_hr(color='1E3A5F', sz='12')

    # ── Sections ────────────────────────────────────────────────
    for section in resume_data.get('sections', []):
        sec_name = section['section_name']
        content  = section.get('content', '')

        sh = doc.add_paragraph()
        sp(sh, 8, 1)
        sr = sh.add_run(sec_name.upper())
        sr.bold = True
        sr.font.size = Pt(10)
        sr.font.color.rgb = ACCENT
        add_hr(color='1E3A5F', sz='8')

        if 'skill' in sec_name.lower():
            skills_text = ' • '.join(
                line.strip() for line in content.split('\n') if line.strip()
            )
            sp_para = doc.add_paragraph(skills_text)
            sp(sp_para, 1, 3)
            for r in sp_para.runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = BLACK
            continue

        entries = _parse_entries(content)
        if not entries:
            pp = doc.add_paragraph(content)
            sp(pp, 1, 2)
            for r in pp.runs:
                r.font.size = Pt(9.5)
            continue

        for headers, bullets in entries:
            for i, h in enumerate(headers):
                left, right = _split_date(h)
                ep = doc.add_paragraph()
                sp(ep, 4 if i == 0 else 0, 1)
                if right:
                    set_right_tab(ep)
                    rl = ep.add_run(left)
                    rl.bold   = (i == 0)
                    rl.italic = (i > 0)
                    rl.font.size      = Pt(9.5)
                    rl.font.color.rgb = BLACK if i == 0 else MGRAY
                    rd = ep.add_run('\t' + right)
                    rd.font.size      = Pt(9)
                    rd.font.color.rgb = MGRAY
                else:
                    rl = ep.add_run(h)
                    rl.bold   = (i == 0)
                    rl.italic = (i > 0)
                    rl.font.size      = Pt(9.5)
                    rl.font.color.rgb = BLACK if i == 0 else MGRAY
            for b in bullets:
                bp = doc.add_paragraph(style='List Bullet')
                sp(bp, 0, 1)
                br = bp.add_run(b)
                br.font.size      = Pt(9.5)
                br.font.color.rgb = BLACK

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def escape_latex(text: str) -> str:
    """Escape special LaTeX characters in plain text."""
    if not text:
        return ""
    for char, rep in [
        ('&', r'\&'), ('%', r'\%'), ('$', r'\$'),
        ('#', r'\#'), ('_', r'\_'),
    ]:
        text = text.replace(char, rep)
    return text

def generate_latex(resume_data: dict) -> str:
    """Fill the ATS LaTeX template with resume data."""
    template_path = ROOT_DIR / 'resume_template.tex'
    with open(template_path, 'r', encoding='utf-8') as f:
        template = f.read()

    sections = {
        s['section_name'].lower(): s.get('content', '')
        for s in resume_data.get('sections', [])
    }

    def lines_of(text, n):
        ls = [l.strip() for l in text.split('\n') if l.strip()]
        return ls[:n] + [''] * max(0, n - len(ls))

    exp_lines = lines_of(sections.get('experience', ''), 6)
    edu_lines = lines_of(sections.get('education', ''), 3)

    vals = {
        'name':               escape_latex(resume_data.get('full_name', '')),
        'job_title':          '',
        'location_link':      'https://maps.google.com',
        'location':           '',
        'phone':              resume_data.get('phone', ''),
        'email':              resume_data.get('email', ''),
        'linkedin_url':       'https://linkedin.com',
        'github_url':         'https://github.com',
        'kaggle_url':         'https://kaggle.com',
        'portfolio_url':      'https://example.com',
        'degree':             escape_latex(edu_lines[0]),
        'graduation_year':    '',
        'university':         escape_latex(edu_lines[1]),
        'cgpa':               '',
        'professional_summary': escape_latex(sections.get('summary', '')),
        'languages':          escape_latex(sections.get('skills', '')),
        'ml_skills':          '',
        'genai_skills':       '',
        'ml_libraries':       '',
        'databases':          '',
        'developer_tools':    '',
        'company_1':          '',
        'duration_1':         '',
        'location_1':         '',
        'experience_point_1': escape_latex(exp_lines[0]),
        'experience_point_2': escape_latex(exp_lines[1]),
        'experience_point_3': escape_latex(exp_lines[2]),
        'company_2':          '',
        'duration_2':         '',
        'location_2':         '',
        'experience_point_4': escape_latex(exp_lines[3]),
        'experience_point_5': escape_latex(exp_lines[4]),
        'experience_point_6': escape_latex(exp_lines[5]),
        'project_1_name':     '',
        'project_1_tech':     '',
        'project_1_demo':     'https://example.com',
        'project_1_github':   'https://github.com',
        'project_1_desc_1':   '',
        'project_1_desc_2':   '',
        'project_2_name':     '',
        'project_2_tech':     '',
        'project_2_demo':     'https://example.com',
        'project_2_github':   'https://github.com',
        'project_2_desc_1':   '',
        'project_2_desc_2':   '',
        'project_3_name':     '',
        'project_3_tech':     '',
        'project_3_desc_1':   '',
        'project_3_desc_2':   '',
        'certifications':     '',
        'achievement_1':      '',
        'achievement_2':      '',
        'achievement_3':      '',
        'spoken_languages':   'English',
        'interests':          '',
    }

    for key, value in vals.items():
        template = template.replace('{{' + key + '}}', value)

    return template

# API Endpoints

# Auth Endpoints
@api_router.post("/auth/register", response_model=Token)
async def register(user_data: UserRegister):
    try:
        # Check if user already exists
        existing_user = await db.users.find_one({"email": user_data.email})
        if existing_user:
            raise HTTPException(status_code=400, detail="Email already registered")
        
        # Create new user
        user_id = str(uuid.uuid4())
        hashed_password = get_password_hash(user_data.password)
        
        user_doc = {
            "id": user_id,
            "email": user_data.email,
            "full_name": user_data.full_name,
            "password_hash": hashed_password,
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        
        await db.users.insert_one(user_doc)
        
        # Create access token
        access_token = create_access_token(data={"sub": user_id, "email": user_data.email})
        
        return {"access_token": access_token, "token_type": "bearer"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Registration error: {e}")
        raise HTTPException(status_code=500, detail="Registration failed")

@api_router.post("/auth/login", response_model=Token)
async def login(user_data: UserLogin):
    try:
        # Find user by email
        user = await db.users.find_one({"email": user_data.email})
        if not user:
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        # Verify password
        if not verify_password(user_data.password, user["password_hash"]):
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        # Create access token
        access_token = create_access_token(data={"sub": user["id"], "email": user["email"]})
        
        return {"access_token": access_token, "token_type": "bearer"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Login error: {e}")
        raise HTTPException(status_code=500, detail="Login failed")

@api_router.post("/auth/google", response_model=Token)
async def google_auth(request: GoogleAuthRequest):
    """Sign in or register with Google OAuth."""
    try:
        from google.oauth2 import id_token
        from google.auth.transport import requests as google_requests

        google_client_id = os.environ.get("GOOGLE_CLIENT_ID")
        if not google_client_id:
            raise HTTPException(status_code=500, detail="Google OAuth not configured. Set GOOGLE_CLIENT_ID in backend .env")

        # Verify the Google ID token
        idinfo = id_token.verify_oauth2_token(
            request.credential,
            google_requests.Request(),
            google_client_id
        )

        google_user_id = idinfo["sub"]
        email = idinfo["email"]
        full_name = idinfo.get("name", email.split("@")[0])

        # Find or create user
        user = await db.users.find_one({"email": email})

        if not user:
            user_id = str(uuid.uuid4())
            user_doc = {
                "id": user_id,
                "email": email,
                "full_name": full_name,
                "google_id": google_user_id,
                "password_hash": None,
                "created_at": datetime.now(timezone.utc).isoformat()
            }
            await db.users.insert_one(user_doc)
        else:
            user_id = user["id"]
            if not user.get("google_id"):
                await db.users.update_one(
                    {"id": user_id},
                    {"$set": {"google_id": google_user_id}}
                )

        access_token = create_access_token(data={"sub": user_id, "email": email})
        return {"access_token": access_token, "token_type": "bearer"}

    except ValueError as e:
        raise HTTPException(status_code=401, detail=f"Invalid Google token: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Google auth error: {e}")
        raise HTTPException(status_code=500, detail="Google authentication failed")

@api_router.get("/auth/me", response_model=User)
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = decode_token(token)
        if not payload:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        user_id = payload.get("sub")
        user = await db.users.find_one({"id": user_id}, {"_id": 0, "password_hash": 0})
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        user["created_at"] = datetime.fromisoformat(user["created_at"])
        return user
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Auth error: {e}")
        raise HTTPException(status_code=401, detail="Authentication failed")

# Helper to get current user from token
async def get_current_user_id(credentials: HTTPAuthorizationCredentials = Depends(security)) -> str:
    payload = decode_token(credentials.credentials)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid token")
    user_id = payload.get("sub")
    if not user_id:
        raise HTTPException(status_code=401, detail="Invalid token")
    return user_id

@api_router.get("/")
async def root():
    return {"message": "CareerArchitect API - AI Resume Builder"}

@api_router.post("/resume/upload")
async def upload_resume(file: UploadFile = File(...), user_id: str = Depends(get_current_user_id)):
    try:
        content = await file.read()
        
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(content)
        elif file.filename.endswith('.docx'):
            text = extract_text_from_docx(content)
        else:
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        sections = parse_resume_sections(text)
        resume = ResumeData(raw_text=text, sections=sections)
        
        doc = resume.model_dump()
        doc['user_id'] = user_id
        doc['created_at'] = doc['created_at'].isoformat()
        await db.resumes.insert_one(doc)
        
        ats_score = calculate_ats_score(text, sections)
        ats_score.resume_id = resume.id
        score_doc = ats_score.model_dump()
        score_doc['user_id'] = user_id
        score_doc['created_at'] = score_doc['created_at'].isoformat()
        await db.ats_scores.insert_one(score_doc)
        
        return {
            "resume_id": resume.id,
            "text": text[:500],
            "sections": [s.model_dump() for s in sections],
            "ats_score": ats_score.model_dump()
        }
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/resume/manual")
async def create_manual_resume(input_data: ManualResumeInput, user_id: str = Depends(get_current_user_id)):
    try:
        sections = [
            ResumeSection(section_name="Summary", content=input_data.summary),
            ResumeSection(section_name="Experience", content=input_data.experience),
            ResumeSection(section_name="Education", content=input_data.education),
            ResumeSection(section_name="Skills", content=input_data.skills)
        ]
        
        raw_text = f"{input_data.full_name}\n{input_data.email}\n{input_data.phone}\n\n"
        raw_text += f"Summary: {input_data.summary}\n\nExperience: {input_data.experience}\n\n"
        raw_text += f"Education: {input_data.education}\n\nSkills: {input_data.skills}"
        
        resume = ResumeData(raw_text=raw_text, sections=sections)
        
        doc = resume.model_dump()
        doc['user_id'] = user_id
        doc['created_at'] = doc['created_at'].isoformat()
        doc['full_name'] = input_data.full_name
        doc['email'] = input_data.email
        doc['phone'] = input_data.phone
        await db.resumes.insert_one(doc)
        
        ats_score = calculate_ats_score(raw_text, sections)
        ats_score.resume_id = resume.id
        score_doc = ats_score.model_dump()
        score_doc['user_id'] = user_id
        score_doc['created_at'] = score_doc['created_at'].isoformat()
        await db.ats_scores.insert_one(score_doc)
        
        return {
            "resume_id": resume.id,
            "sections": [s.model_dump() for s in sections],
            "ats_score": ats_score.model_dump()
        }
    except Exception as e:
        logger.error(f"Manual resume error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/resume/enhance")
async def enhance_resume(request: EnhanceRequest, user_id: str = Depends(get_current_user_id)):
    try:
        resume = await db.resumes.find_one({"id": request.resume_id}, {"_id": 0})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Verify ownership
        if resume.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Unauthorized")
        
        enhanced_text = ""
        enhancement_type = request.enhancement_type
        
        if enhancement_type == "openai":
            enhanced_text = await enhance_with_openai(resume['raw_text'])
        elif enhancement_type == "gemini":
            enhanced_text = await enhance_with_gemini(resume['raw_text'])
        else:
            openai_enhanced = await enhance_with_openai(resume['raw_text'])
            enhanced_text = await enhance_with_gemini(openai_enhanced)
        
        enhanced_sections = parse_resume_sections(enhanced_text)
        
        enhanced_resume = EnhancedResume(
            original_resume_id=request.resume_id,
            enhanced_text=enhanced_text,
            enhanced_sections=enhanced_sections,
            enhancement_type=enhancement_type
        )
        
        doc = enhanced_resume.model_dump()
        doc['user_id'] = user_id
        doc['created_at'] = doc['created_at'].isoformat()
        await db.enhanced_resumes.insert_one(doc)
        
        new_ats_score = calculate_ats_score(enhanced_text, enhanced_sections)
        new_ats_score.resume_id = enhanced_resume.id
        score_doc = new_ats_score.model_dump()
        score_doc['user_id'] = user_id
        score_doc['created_at'] = score_doc['created_at'].isoformat()
        await db.ats_scores.insert_one(score_doc)
        
        return {
            "enhanced_resume_id": enhanced_resume.id,
            "enhanced_text": enhanced_text,
            "enhanced_sections": [s.model_dump() for s in enhanced_sections],
            "new_ats_score": new_ats_score.model_dump()
        }
    except Exception as e:
        logger.error(f"Enhancement error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/resume/{resume_id}")
async def get_resume(resume_id: str, user_id: str = Depends(get_current_user_id)):
    try:
        resume = await db.resumes.find_one({"id": resume_id}, {"_id": 0})
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        
        # Verify ownership
        if resume.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Unauthorized")
        
        ats_score = await db.ats_scores.find_one({"resume_id": resume_id}, {"_id": 0})
        
        return {
            "resume": resume,
            "ats_score": ats_score
        }
    except Exception as e:
        logger.error(f"Get resume error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@api_router.post("/resume/generate/{resume_id}")
async def generate_resume(resume_id: str, format: str = "pdf", user_id: str = Depends(get_current_user_id)):
    try:
        resume = await db.resumes.find_one({"id": resume_id}, {"_id": 0})
        if not resume:
            enhanced = await db.enhanced_resumes.find_one({"id": resume_id}, {"_id": 0})
            if not enhanced:
                raise HTTPException(status_code=404, detail="Resume not found")
            resume = enhanced
        
        # Verify ownership
        if resume.get("user_id") != user_id:
            raise HTTPException(status_code=403, detail="Unauthorized")
        
        if format == "pdf":
            pdf_bytes = generate_pdf(resume)
            return {"file_data": pdf_bytes.hex(), "format": "pdf"}
        elif format == "docx":
            docx_bytes = generate_docx(resume)
            return {"file_data": docx_bytes.hex(), "format": "docx"}
        elif format == "latex":
            latex_str = generate_latex(resume)
            latex_bytes = latex_str.encode('utf-8')
            return {"file_data": latex_bytes.hex(), "format": "latex"}
        else:
            raise HTTPException(status_code=400, detail="Format must be 'pdf', 'docx', or 'latex'")
    except Exception as e:
        logger.error(f"Generate error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

app.include_router(api_router)

@app.get("/")
async def app_root():
    return {
        "message": "CareerArchitect API - AI Resume Builder",
        "docs": "/docs",
        "api": "/api/"
    }

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()